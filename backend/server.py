from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import asyncio
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import resend
from bson import ObjectId
import base64
import io

# Document generation imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET')
if not JWT_SECRET:
    raise ValueError("JWT_SECRET environment variable must be set")
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# Resend Configuration
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'onboarding@resend.dev')
if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

# Create the main app
app = FastAPI(title="WBS Transcript and Recommendation Tracker API")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Upload directory
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# ==================== MODELS ====================

class UserBase(BaseModel):
    email: EmailStr
    full_name: str
    role: str = "student"  # student, staff, admin

class UserCreate(BaseModel):
    email: EmailStr
    full_name: str
    password: str
    role: str = "student"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    full_name: str
    role: str
    created_at: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class TranscriptRequestCreate(BaseModel):
    first_name: str
    middle_name: Optional[str] = ""
    last_name: str
    school_id: str
    enrollment_status: str  # enrolled, graduate, withdrawn
    academic_years: List[dict]  # List of {"from_year": "2015", "to_year": "2020"}
    wolmers_email: str
    personal_email: EmailStr
    phone_number: str
    reason: str
    needed_by_date: str
    collection_method: str  # pickup, emailed, delivery
    delivery_address: Optional[str] = ""  # Required if collection_method is 'delivery'
    institution_name: Optional[str] = ""
    institution_address: Optional[str] = ""
    institution_phone: Optional[str] = ""
    institution_email: Optional[str] = ""

class PasswordResetRequest(BaseModel):
    email: EmailStr

class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

class TranscriptRequestUpdate(BaseModel):
    status: Optional[str] = None
    assigned_staff_id: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    note: Optional[str] = None  # Note for status changes

class TranscriptRequestResponse(BaseModel):
    id: str
    student_id: str
    student_name: str
    student_email: str
    first_name: str
    middle_name: str
    last_name: str
    school_id: str
    enrollment_status: str
    academic_years: List[dict] = []  # List of {"from_year": "2015", "to_year": "2020"}
    academic_year: str = ""  # Legacy field for backward compatibility
    wolmers_email: str
    personal_email: str
    phone_number: str
    reason: str
    needed_by_date: str
    collection_method: str
    delivery_address: str = ""
    institution_name: str = ""
    institution_address: str
    institution_phone: str
    institution_email: str
    status: str
    assigned_staff_id: Optional[str] = None
    assigned_staff_name: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    documents: List[dict] = []
    timeline: List[dict] = []
    created_at: str
    updated_at: str

class NotificationResponse(BaseModel):
    id: str
    user_id: str
    title: str
    message: str
    type: str
    read: bool
    request_id: Optional[str] = None
    created_at: str

class StaffCreateByAdmin(BaseModel):
    email: EmailStr
    full_name: str
    password: str
    role: str  # staff or admin

class AnalyticsResponse(BaseModel):
    total_requests: int
    pending_requests: int
    in_progress_requests: int
    processing_requests: int
    ready_requests: int
    completed_requests: int
    rejected_requests: int
    overdue_requests: int
    requests_by_month: List[dict]
    requests_by_enrollment: List[dict]
    requests_by_collection_method: List[dict]
    staff_workload: List[dict]
    overdue_by_days: List[dict]
    # Recommendation letter stats
    total_recommendation_requests: int = 0
    pending_recommendation_requests: int = 0
    in_progress_recommendation_requests: int = 0
    completed_recommendation_requests: int = 0
    rejected_recommendation_requests: int = 0
    overdue_recommendation_requests: int = 0
    # Collection method breakdown for recommendations
    recommendations_by_collection_method: List[dict] = []
    # Overdue breakdown
    overdue_transcripts_by_days: List[dict] = []
    overdue_recommendations_by_days: List[dict] = []

# ==================== RECOMMENDATION LETTER MODELS ====================

class RecommendationRequestCreate(BaseModel):
    first_name: str
    middle_name: Optional[str] = ""
    last_name: str
    email: EmailStr
    phone_number: str
    address: str
    years_attended: List[dict]  # List of {"from_year": "2015", "to_year": "2020"}
    last_form_class: str  # e.g., "6th Form" or "Upper 6"
    co_curricular_activities: Optional[str] = ""  # Positions of responsibility and activities
    institution_name: str
    institution_address: str
    directed_to: Optional[str] = ""  # Whom should the letter be directed to
    program_name: str
    needed_by_date: str
    collection_method: str  # pickup, emailed, delivery
    delivery_address: Optional[str] = ""  # Required if collection_method is 'delivery'

class RecommendationRequestUpdate(BaseModel):
    status: Optional[str] = None
    assigned_staff_id: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    co_curricular_activities: Optional[str] = None
    note: Optional[str] = None  # Note for status changes

class RecommendationRequestResponse(BaseModel):
    id: str
    student_id: str
    student_name: str
    student_email: str
    first_name: str
    middle_name: str
    last_name: str
    email: str
    phone_number: str
    address: str
    years_attended: List[dict] = []  # List of {"from_year": "2015", "to_year": "2020"}
    years_attended_str: str = ""  # Legacy string format for backward compatibility
    last_form_class: str
    co_curricular_activities: str = ""
    institution_name: str
    institution_address: str
    directed_to: str
    program_name: str
    needed_by_date: str
    collection_method: str
    delivery_address: str = ""
    status: str
    assigned_staff_id: Optional[str] = None
    assigned_staff_name: Optional[str] = None
    rejection_reason: Optional[str] = None
    staff_notes: Optional[str] = None
    documents: List[dict] = []
    timeline: List[dict] = []
    created_at: str
    updated_at: str

class StudentRecommendationUpdate(BaseModel):
    first_name: Optional[str] = None
    middle_name: Optional[str] = None
    last_name: Optional[str] = None
    email: Optional[str] = None
    phone_number: Optional[str] = None
    address: Optional[str] = None
    years_attended: Optional[List[dict]] = None
    last_form_class: Optional[str] = None
    co_curricular_activities: Optional[str] = None
    institution_name: Optional[str] = None
    institution_address: Optional[str] = None
    directed_to: Optional[str] = None
    program_name: Optional[str] = None
    needed_by_date: Optional[str] = None
    collection_method: Optional[str] = None
    delivery_address: Optional[str] = None

# ==================== HELPER FUNCTIONS ====================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def decode_token(token: str) -> dict:
    try:
        return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    payload = decode_token(token)
    user = await db.users.find_one({"id": payload["sub"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user

async def require_role(roles: List[str]):
    async def role_checker(user: dict = Depends(get_current_user)):
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return user
    return role_checker

async def send_email_notification(to_email: str, subject: str, html_content: str):
    if not RESEND_API_KEY:
        logger.warning("Resend API key not configured, skipping email")
        return None
    try:
        params = {
            "from": SENDER_EMAIL,
            "to": [to_email],
            "subject": subject,
            "html": html_content
        }
        result = await asyncio.to_thread(resend.Emails.send, params)
        logger.info(f"Email sent to {to_email}")
        return result
    except Exception as e:
        logger.error(f"Failed to send email: {str(e)}")
        return None

async def create_notification(user_id: str, title: str, message: str, notif_type: str, request_id: str = None):
    notification = {
        "id": str(uuid.uuid4()),
        "user_id": user_id,
        "title": title,
        "message": message,
        "type": notif_type,
        "read": False,
        "request_id": request_id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(notification)
    return notification

async def check_and_notify_overdue_requests():
    """Check for overdue requests and notify admins"""
    now = datetime.now(timezone.utc)
    today_str = now.strftime("%Y-%m-%d")
    
    # Find overdue requests that haven't been notified today
    overdue_requests = await db.transcript_requests.find({
        "needed_by_date": {"$lt": today_str},
        "status": {"$nin": ["Completed", "Rejected"]},
        "$or": [
            {"overdue_notified_date": {"$exists": False}},
            {"overdue_notified_date": {"$ne": today_str}}
        ]
    }).to_list(None)
    
    if not overdue_requests:
        return
    
    # Get all admin users
    admins = await db.users.find({"role": "admin"}, {"_id": 0, "id": 1}).to_list(None)
    
    for req in overdue_requests:
        try:
            needed_date = datetime.strptime(req["needed_by_date"], "%Y-%m-%d")
            days_overdue = (now.replace(tzinfo=None) - needed_date).days
            
            student_name = f"{req.get('first_name', '')} {req.get('last_name', '')}"
            title = "⚠️ Overdue Transcript Request"
            message = f"Request from {student_name} is {days_overdue} day(s) overdue. Needed by: {req['needed_by_date']}"
            
            # Notify all admins
            for admin in admins:
                await create_notification(admin["id"], title, message, "overdue", req["id"])
            
            # Mark as notified today
            await db.transcript_requests.update_one(
                {"id": req["id"]},
                {"$set": {"overdue_notified_date": today_str}}
            )
        except Exception as e:
            print(f"Error notifying overdue request: {e}")

def normalize_recommendation_data(request_data: dict) -> dict:
    """Normalize recommendation request data for backward compatibility"""
    # Make a copy to avoid modifying original
    data = dict(request_data)
    
    # Handle years_attended field migration from string to list
    years_attended = data.get("years_attended", None)
    if years_attended is None:
        data["years_attended"] = []
        data["years_attended_str"] = ""
    elif isinstance(years_attended, str):
        # Convert old string format to new list format
        if years_attended and years_attended.strip() != "":
            # Handle formats like "2015-2020" or "2015-2020, 2021-2022"
            years_list = []
            for year_range in years_attended.split(", "):
                year_range = year_range.strip()
                if "-" in year_range:
                    parts = year_range.split("-", 1)
                    if len(parts) == 2:
                        years_list.append({"from_year": parts[0].strip(), "to_year": parts[1].strip()})
            data["years_attended"] = years_list
            data["years_attended_str"] = years_attended
        else:
            data["years_attended"] = []
            data["years_attended_str"] = ""
    elif isinstance(years_attended, list):
        # Already in new format, create string version for backward compatibility
        years_str = ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in years_attended if isinstance(y, dict)])
        data["years_attended_str"] = years_str
    else:
        data["years_attended"] = []
        data["years_attended_str"] = ""
    
    # Ensure all required fields have default values
    if "co_curricular_activities" not in data or data.get("co_curricular_activities") is None:
        data["co_curricular_activities"] = ""
    if "delivery_address" not in data or data.get("delivery_address") is None:
        data["delivery_address"] = ""
    if "years_attended_str" not in data:
        data["years_attended_str"] = ""
    
    return data

def normalize_transcript_data(request_data: dict) -> dict:
    """Normalize transcript request data for backward compatibility"""
    # Make a copy to avoid modifying original
    data = dict(request_data)
    
    # Handle academic_years field migration from string to list
    academic_years = data.get("academic_years", None)
    if academic_years is None or (isinstance(academic_years, str)):
        # Convert old string format to new list format
        legacy_year = data.get("academic_year", "")
        if legacy_year and isinstance(legacy_year, str) and legacy_year.strip() != "":
            # Use the legacy academic_year field
            years_list = []
            for year_range in legacy_year.split(", "):
                year_range = year_range.strip()
                if "-" in year_range:
                    parts = year_range.split("-", 1)
                    if len(parts) == 2:
                        years_list.append({"from_year": parts[0].strip(), "to_year": parts[1].strip()})
            data["academic_years"] = years_list
        else:
            data["academic_years"] = []
    elif isinstance(academic_years, list):
        # Already in new format, create string version for backward compatibility
        years_str = ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in academic_years if isinstance(y, dict)])
        data["academic_year"] = years_str
    else:
        data["academic_years"] = []
    
    # Ensure academic_year legacy field exists
    if "academic_year" not in data or data.get("academic_year") is None:
        data["academic_year"] = ""
    
    # Ensure all required fields have default values
    if "delivery_address" not in data or data.get("delivery_address") is None:
        data["delivery_address"] = ""
    
    return data

async def notify_status_change(request_data: dict, old_status: str, new_status: str):
    student = await db.users.find_one({"id": request_data["student_id"]}, {"_id": 0})
    if not student:
        return
    
    title = f"Request Status Updated"
    message = f"Your transcript request has been updated from '{old_status}' to '{new_status}'."
    
    # Create in-app notification
    await create_notification(student["id"], title, message, "status_update", request_data["id"])
    
    # Send email notification
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background-color: #800000; color: white; padding: 20px; text-align: center;">
            <h1 style="margin: 0;">Wolmer's Boys' School</h1>
            <p style="margin: 5px 0;">Transcript Tracker</p>
        </div>
        <div style="padding: 20px; background-color: #f5f5f5;">
            <h2 style="color: #800000;">Request Status Update</h2>
            <p>Dear {student['full_name']},</p>
            <p>Your transcript request status has been updated:</p>
            <div style="background-color: white; padding: 15px; border-radius: 5px; margin: 15px 0;">
                <p><strong>Previous Status:</strong> {old_status}</p>
                <p><strong>New Status:</strong> <span style="color: #800000; font-weight: bold;">{new_status}</span></p>
            </div>
            <p>Log in to view more details about your request.</p>
            <p style="color: #666; font-size: 12px; margin-top: 30px;">
                This is an automated message from Wolmer's Boys' School Transcript Tracker.
            </p>
        </div>
    </div>
    """
    await send_email_notification(student["email"], f"Transcript Request: {new_status}", html_content)

# ==================== AUTH ROUTES ====================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    # Check if email already exists
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Only students can self-register
    if user_data.role != "student":
        raise HTTPException(status_code=400, detail="Only students can self-register")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "id": user_id,
        "email": user_data.email,
        "full_name": user_data.full_name,
        "password_hash": hash_password(user_data.password),
        "role": "student",
        "created_at": now,
        "updated_at": now
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id, user_data.email, "student")
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user_id,
            email=user_data.email,
            full_name=user_data.full_name,
            role="student",
            created_at=now
        )
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    
    token = create_token(user["id"], user["email"], user["role"])
    
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user["id"],
            email=user["email"],
            full_name=user["full_name"],
            role=user["role"],
            created_at=user["created_at"]
        )
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(user: dict = Depends(get_current_user)):
    return UserResponse(
        id=user["id"],
        email=user["email"],
        full_name=user["full_name"],
        role=user["role"],
        created_at=user["created_at"]
    )

# ==================== PASSWORD RESET ====================

@api_router.post("/auth/forgot-password")
async def forgot_password(request: PasswordResetRequest):
    user = await db.users.find_one({"email": request.email}, {"_id": 0})
    if not user:
        # Don't reveal if email exists or not for security
        return {"message": "If an account with this email exists, a password reset link has been sent."}
    
    # Generate reset token (valid for 1 hour)
    reset_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Store reset token in database
    await db.password_resets.delete_many({"email": request.email})  # Remove old tokens
    await db.password_resets.insert_one({
        "token": reset_token,
        "email": request.email,
        "user_id": user["id"],
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    # Send email with reset link
    frontend_url = os.environ.get('FRONTEND_URL', 'http://localhost:3000')
    reset_link = f"{frontend_url}/reset-password?token={reset_token}"
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background-color: #800000; color: white; padding: 20px; text-align: center;">
            <h1 style="margin: 0;">Wolmer's Boys' School</h1>
            <p style="margin: 5px 0;">Transcript Tracker</p>
        </div>
        <div style="padding: 20px; background-color: #f5f5f5;">
            <h2 style="color: #800000;">Password Reset Request</h2>
            <p>Dear {user['full_name']},</p>
            <p>We received a request to reset your password. Click the button below to create a new password:</p>
            <div style="text-align: center; margin: 30px 0;">
                <a href="{reset_link}" style="background-color: #800000; color: white; padding: 12px 30px; text-decoration: none; border-radius: 5px; display: inline-block;">Reset Password</a>
            </div>
            <p style="color: #666; font-size: 14px;">This link will expire in 1 hour.</p>
            <p style="color: #666; font-size: 14px;">If you didn't request this, please ignore this email.</p>
            <p style="color: #999; font-size: 12px; margin-top: 30px;">
                This is an automated message from Wolmer's Boys' School Transcript Tracker.
            </p>
        </div>
    </div>
    """
    await send_email_notification(request.email, "Password Reset Request", html_content)
    
    # For development: log the token
    logger.info(f"Password reset token for {request.email}: {reset_token}")
    
    return {"message": "If an account with this email exists, a password reset link has been sent.", "token": reset_token}

@api_router.post("/auth/reset-password")
async def reset_password(request: PasswordResetConfirm):
    # Find the reset token
    reset_record = await db.password_resets.find_one({"token": request.token}, {"_id": 0})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Check if token is expired
    expires_at = datetime.fromisoformat(reset_record["expires_at"])
    if datetime.now(timezone.utc) > expires_at:
        await db.password_resets.delete_one({"token": request.token})
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    # Update user's password
    new_password_hash = hash_password(request.new_password)
    await db.users.update_one(
        {"id": reset_record["user_id"]},
        {"$set": {"password_hash": new_password_hash, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Delete the used token
    await db.password_resets.delete_one({"token": request.token})
    
    return {"message": "Password has been reset successfully"}

@api_router.get("/auth/verify-reset-token/{token}")
async def verify_reset_token(token: str):
    reset_record = await db.password_resets.find_one({"token": token}, {"_id": 0})
    
    if not reset_record:
        raise HTTPException(status_code=400, detail="Invalid reset token")
    
    expires_at = datetime.fromisoformat(reset_record["expires_at"])
    if datetime.now(timezone.utc) > expires_at:
        await db.password_resets.delete_one({"token": token})
        raise HTTPException(status_code=400, detail="Reset token has expired")
    
    return {"valid": True, "email": reset_record["email"]}

# ==================== USER MANAGEMENT (ADMIN) ====================

@api_router.post("/admin/users", response_model=UserResponse)
async def create_user_by_admin(user_data: StaffCreateByAdmin, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can create users")
    
    existing = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "id": user_id,
        "email": user_data.email,
        "full_name": user_data.full_name,
        "password_hash": hash_password(user_data.password),
        "role": user_data.role,
        "created_at": now,
        "updated_at": now
    }
    
    await db.users.insert_one(user_doc)
    
    return UserResponse(
        id=user_id,
        email=user_data.email,
        full_name=user_data.full_name,
        role=user_data.role,
        created_at=now
    )

@api_router.get("/admin/users", response_model=List[UserResponse])
async def get_all_users(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view all users")
    
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return [UserResponse(**u) for u in users]

@api_router.get("/admin/staff", response_model=List[UserResponse])
async def get_staff_members(current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    staff = await db.users.find({"role": "staff"}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return [UserResponse(**s) for s in staff]

@api_router.delete("/admin/users/{user_id}")
async def delete_user(user_id: str, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can delete users")
    
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Cannot delete yourself")
    
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User deleted successfully"}

class AdminResetPassword(BaseModel):
    new_password: str

@api_router.post("/admin/users/{user_id}/reset-password")
async def admin_reset_user_password(user_id: str, data: AdminResetPassword, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can reset passwords")
    
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Cannot reset your own password this way. Use the forgot password feature.")
    
    # Find the user
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Update password
    new_password_hash = hash_password(data.new_password)
    await db.users.update_one(
        {"id": user_id},
        {"$set": {"password_hash": new_password_hash, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    # Send notification email to user
    html_content = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background-color: #800000; color: white; padding: 20px; text-align: center;">
            <h1 style="margin: 0;">Wolmer's Boys' School</h1>
            <p style="margin: 5px 0;">Transcript Tracker</p>
        </div>
        <div style="padding: 20px; background-color: #f5f5f5;">
            <h2 style="color: #800000;">Password Reset by Administrator</h2>
            <p>Dear {user['full_name']},</p>
            <p>Your password has been reset by an administrator.</p>
            <p>If you did not request this change, please contact the administrator immediately.</p>
            <p style="color: #999; font-size: 12px; margin-top: 30px;">
                This is an automated message from Wolmer's Boys' School Transcript Tracker.
            </p>
        </div>
    </div>
    """
    await send_email_notification(user["email"], "Your Password Has Been Reset", html_content)
    
    return {"message": f"Password reset successfully for {user['full_name']}"}

# ==================== TRANSCRIPT REQUESTS ====================

@api_router.post("/requests", response_model=TranscriptRequestResponse)
async def create_transcript_request(request_data: TranscriptRequestCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can create transcript requests")
    
    request_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request submitted",
        "updated_by": current_user["full_name"]
    }
    
    # Format academic years for display (backward compatibility)
    academic_years_str = ", ".join([f"{y['from_year']}-{y['to_year']}" for y in request_data.academic_years]) if request_data.academic_years else ""
    
    doc = {
        "id": request_id,
        "student_id": current_user["id"],
        "student_name": current_user["full_name"],
        "student_email": current_user["email"],
        "first_name": request_data.first_name,
        "middle_name": request_data.middle_name or "",
        "last_name": request_data.last_name,
        "school_id": request_data.school_id,
        "enrollment_status": request_data.enrollment_status,
        "academic_years": request_data.academic_years,
        "academic_year": academic_years_str,  # Legacy field
        "wolmers_email": request_data.wolmers_email,
        "personal_email": request_data.personal_email,
        "phone_number": request_data.phone_number,
        "reason": request_data.reason,
        "needed_by_date": request_data.needed_by_date,
        "collection_method": request_data.collection_method,
        "delivery_address": request_data.delivery_address or "",
        "institution_name": request_data.institution_name or "",
        "institution_address": request_data.institution_address or "",
        "institution_phone": request_data.institution_phone or "",
        "institution_email": request_data.institution_email or "",
        "status": "Pending",
        "assigned_staff_id": None,
        "assigned_staff_name": None,
        "rejection_reason": None,
        "staff_notes": None,
        "documents": [],
        "timeline": [timeline_entry],
        "created_at": now,
        "updated_at": now
    }
    
    await db.transcript_requests.insert_one(doc)
    
    # Notify admins
    admins = await db.users.find({"role": "admin"}, {"_id": 0}).to_list(100)
    for admin in admins:
        await create_notification(
            admin["id"],
            "New Transcript Request",
            f"New request from {current_user['full_name']}",
            "new_request",
            request_id
        )
    
    return TranscriptRequestResponse(**doc)

@api_router.get("/requests", response_model=List[TranscriptRequestResponse])
async def get_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] == "student":
        # Students can only see their own requests
        requests = await db.transcript_requests.find(
            {"student_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    elif current_user["role"] == "staff":
        # Staff can see assigned requests
        requests = await db.transcript_requests.find(
            {"assigned_staff_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    else:
        # Admin can see all requests
        requests = await db.transcript_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Normalize data for backward compatibility
    normalized_requests = [normalize_transcript_data(r) for r in requests]
    return [TranscriptRequestResponse(**r) for r in normalized_requests]

@api_router.get("/requests/all", response_model=List[TranscriptRequestResponse])
async def get_all_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    requests = await db.transcript_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    # Normalize data for backward compatibility
    normalized_requests = [normalize_transcript_data(r) for r in requests]
    return [TranscriptRequestResponse(**r) for r in normalized_requests]

@api_router.get("/requests/{request_id}", response_model=TranscriptRequestResponse)
async def get_request(request_id: str, current_user: dict = Depends(get_current_user)):
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check permissions
    if current_user["role"] == "student" and request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only view your own requests")
    
    # Normalize data for backward compatibility
    normalized_request = normalize_transcript_data(request_doc)
    return TranscriptRequestResponse(**normalized_request)

class StudentRequestUpdate(BaseModel):
    first_name: Optional[str] = None
    middle_name: Optional[str] = None
    last_name: Optional[str] = None
    school_id: Optional[str] = None
    enrollment_status: Optional[str] = None
    academic_year: Optional[str] = None
    wolmers_email: Optional[str] = None
    personal_email: Optional[str] = None
    phone_number: Optional[str] = None
    reason: Optional[str] = None
    needed_by_date: Optional[str] = None
    collection_method: Optional[str] = None
    institution_name: Optional[str] = None
    institution_address: Optional[str] = None
    institution_phone: Optional[str] = None
    institution_email: Optional[str] = None

@api_router.put("/requests/{request_id}/edit", response_model=TranscriptRequestResponse)
async def student_edit_request(request_id: str, update_data: StudentRequestUpdate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can edit their own requests")
    
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check ownership
    if request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only edit your own requests")
    
    # Check status - only pending requests can be edited
    if request_doc["status"] != "Pending":
        raise HTTPException(
            status_code=400, 
            detail=f"This request cannot be edited because its status is '{request_doc['status']}'. Only pending requests can be modified."
        )
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    # Update only provided fields
    update_fields = update_data.dict(exclude_unset=True)
    for field, value in update_fields.items():
        if value is not None:
            updates[field] = value
    
    # Add timeline entry for edit
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request details updated by student",
        "updated_by": current_user["full_name"]
    }
    
    await db.transcript_requests.update_one(
        {"id": request_id},
        {
            "$set": updates,
            "$push": {"timeline": timeline_entry}
        }
    )
    
    updated_request = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    return TranscriptRequestResponse(**updated_request)

@api_router.patch("/requests/{request_id}", response_model=TranscriptRequestResponse)
async def update_request(request_id: str, update_data: TranscriptRequestUpdate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] == "student":
        raise HTTPException(status_code=403, detail="Students cannot update request status")
    
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    old_status = request_doc["status"]
    
    if update_data.status:
        updates["status"] = update_data.status
        # Use the provided note or create a default one
        note_text = update_data.note if update_data.note else f"Status changed to {update_data.status}"
        timeline_entry = {
            "status": update_data.status,
            "timestamp": now,
            "note": note_text,
            "updated_by": current_user["full_name"]
        }
        await db.transcript_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.assigned_staff_id:
        staff = await db.users.find_one({"id": update_data.assigned_staff_id}, {"_id": 0})
        if staff:
            updates["assigned_staff_id"] = update_data.assigned_staff_id
            updates["assigned_staff_name"] = staff["full_name"]
            # Notify staff
            await create_notification(
                staff["id"],
                "New Assignment",
                f"You have been assigned a transcript request",
                "assignment",
                request_id
            )
    
    if update_data.rejection_reason:
        updates["rejection_reason"] = update_data.rejection_reason
        updates["status"] = "Rejected"
        timeline_entry = {
            "status": "Rejected",
            "timestamp": now,
            "note": f"Request rejected: {update_data.rejection_reason}",
            "updated_by": current_user["full_name"]
        }
        await db.transcript_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.staff_notes:
        updates["staff_notes"] = update_data.staff_notes
    
    await db.transcript_requests.update_one({"id": request_id}, {"$set": updates})
    
    # Notify student of status change
    if update_data.status and update_data.status != old_status:
        updated_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
        await notify_status_change(updated_doc, old_status, update_data.status)
    
    updated_request = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    return TranscriptRequestResponse(**updated_request)

# ==================== FILE UPLOAD ====================

@api_router.post("/requests/{request_id}/documents")
async def upload_document(request_id: str, file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Only staff and admin can upload documents")
    
    request_doc = await db.transcript_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Validate file type
    allowed_types = [
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "image/jpeg",
        "image/png",
        "image/gif"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not allowed")
    
    # Save file
    file_id = str(uuid.uuid4())
    file_ext = Path(file.filename).suffix
    file_path = UPLOAD_DIR / f"{file_id}{file_ext}"
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    now = datetime.now(timezone.utc).isoformat()
    doc_entry = {
        "id": file_id,
        "filename": file.filename,
        "content_type": file.content_type,
        "path": str(file_path),
        "uploaded_by": current_user["full_name"],
        "uploaded_at": now
    }
    
    await db.transcript_requests.update_one(
        {"id": request_id},
        {
            "$push": {"documents": doc_entry},
            "$set": {"updated_at": now}
        }
    )
    
    # Add timeline entry
    timeline_entry = {
        "status": request_doc["status"],
        "timestamp": now,
        "note": f"Document uploaded: {file.filename}",
        "updated_by": current_user["full_name"]
    }
    await db.transcript_requests.update_one(
        {"id": request_id},
        {"$push": {"timeline": timeline_entry}}
    )
    
    # Notify student
    await create_notification(
        request_doc["student_id"],
        "Document Uploaded",
        f"A document has been uploaded to your transcript request",
        "document",
        request_id
    )
    
    return {"message": "Document uploaded successfully", "document": doc_entry}

@api_router.get("/documents/{document_id}")
async def get_document(document_id: str, current_user: dict = Depends(get_current_user)):
    # Find the request containing this document
    request_doc = await db.transcript_requests.find_one(
        {"documents.id": document_id},
        {"_id": 0}
    )
    
    if not request_doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check permissions
    if current_user["role"] == "student" and request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Find the document
    doc = next((d for d in request_doc["documents"] if d["id"] == document_id), None)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = Path(doc["path"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on server")
    
    with open(file_path, "rb") as f:
        content = base64.b64encode(f.read()).decode('utf-8')
    
    return {
        "filename": doc["filename"],
        "content_type": doc["content_type"],
        "content": content
    }

# ==================== RECOMMENDATION LETTER REQUESTS ====================

@api_router.post("/recommendations", response_model=RecommendationRequestResponse)
async def create_recommendation_request(request_data: RecommendationRequestCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can create recommendation letter requests")
    
    request_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request submitted",
        "updated_by": current_user["full_name"]
    }
    
    # Format years attended for display (backward compatibility)
    years_attended_str = ", ".join([f"{y['from_year']}-{y['to_year']}" for y in request_data.years_attended]) if request_data.years_attended else ""
    
    doc = {
        "id": request_id,
        "student_id": current_user["id"],
        "student_name": current_user["full_name"],
        "student_email": current_user["email"],
        "first_name": request_data.first_name,
        "middle_name": request_data.middle_name or "",
        "last_name": request_data.last_name,
        "email": request_data.email,
        "phone_number": request_data.phone_number,
        "address": request_data.address,
        "years_attended": request_data.years_attended,
        "years_attended_str": years_attended_str,  # Legacy field
        "last_form_class": request_data.last_form_class,
        "co_curricular_activities": request_data.co_curricular_activities or "",
        "institution_name": request_data.institution_name,
        "institution_address": request_data.institution_address,
        "directed_to": request_data.directed_to or "",
        "program_name": request_data.program_name,
        "needed_by_date": request_data.needed_by_date,
        "collection_method": request_data.collection_method,
        "delivery_address": request_data.delivery_address or "",
        "status": "Pending",
        "assigned_staff_id": None,
        "assigned_staff_name": None,
        "rejection_reason": None,
        "staff_notes": None,
        "documents": [],
        "timeline": [timeline_entry],
        "created_at": now,
        "updated_at": now
    }
    
    await db.recommendation_requests.insert_one(doc)
    
    # Notify admins
    admins = await db.users.find({"role": "admin"}, {"_id": 0}).to_list(100)
    for admin in admins:
        await create_notification(
            admin["id"],
            "New Recommendation Letter Request",
            f"New recommendation letter request from {current_user['full_name']}",
            "new_recommendation",
            request_id
        )
    
    return RecommendationRequestResponse(**doc)

@api_router.get("/recommendations", response_model=List[RecommendationRequestResponse])
async def get_recommendation_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] == "student":
        # Students can only see their own requests
        requests = await db.recommendation_requests.find(
            {"student_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    elif current_user["role"] == "staff":
        # Staff can see assigned requests
        requests = await db.recommendation_requests.find(
            {"assigned_staff_id": current_user["id"]},
            {"_id": 0}
        ).sort("created_at", -1).to_list(1000)
    else:
        # Admin can see all requests
        requests = await db.recommendation_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Normalize data for backward compatibility
    normalized_requests = [normalize_recommendation_data(r) for r in requests]
    return [RecommendationRequestResponse(**r) for r in normalized_requests]

@api_router.get("/recommendations/all", response_model=List[RecommendationRequestResponse])
async def get_all_recommendation_requests(current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    requests = await db.recommendation_requests.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    # Normalize data for backward compatibility
    normalized_requests = [normalize_recommendation_data(r) for r in requests]
    return [RecommendationRequestResponse(**r) for r in normalized_requests]

@api_router.get("/recommendations/{request_id}", response_model=RecommendationRequestResponse)
async def get_recommendation_request(request_id: str, current_user: dict = Depends(get_current_user)):
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check permissions
    if current_user["role"] == "student" and request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only view your own requests")
    
    # Normalize data for backward compatibility
    normalized_request = normalize_recommendation_data(request_doc)
    return RecommendationRequestResponse(**normalized_request)

@api_router.put("/recommendations/{request_id}/edit", response_model=RecommendationRequestResponse)
async def student_edit_recommendation(request_id: str, update_data: StudentRecommendationUpdate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "student":
        raise HTTPException(status_code=403, detail="Only students can edit their own requests")
    
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check ownership
    if request_doc["student_id"] != current_user["id"]:
        raise HTTPException(status_code=403, detail="You can only edit your own requests")
    
    # Check status - only pending requests can be edited
    if request_doc["status"] != "Pending":
        raise HTTPException(
            status_code=400, 
            detail=f"This request cannot be edited because its status is '{request_doc['status']}'. Only pending requests can be modified."
        )
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    # Update only provided fields
    update_fields = update_data.dict(exclude_unset=True)
    for field, value in update_fields.items():
        if value is not None:
            updates[field] = value
    
    # Add timeline entry for edit
    timeline_entry = {
        "status": "Pending",
        "timestamp": now,
        "note": "Request details updated by student",
        "updated_by": current_user["full_name"]
    }
    
    await db.recommendation_requests.update_one(
        {"id": request_id},
        {
            "$set": updates,
            "$push": {"timeline": timeline_entry}
        }
    )
    
    updated_request = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    normalized_request = normalize_recommendation_data(updated_request)
    return RecommendationRequestResponse(**normalized_request)

@api_router.patch("/recommendations/{request_id}", response_model=RecommendationRequestResponse)
async def update_recommendation_request(request_id: str, update_data: RecommendationRequestUpdate, current_user: dict = Depends(get_current_user)):
    # Allow students to update their own pending recommendations
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Check permissions
    if current_user["role"] == "student":
        # Students can only update their own requests and only if status is Pending
        if request_doc["student_id"] != current_user["id"]:
            raise HTTPException(status_code=403, detail="You can only update your own requests")
        if request_doc["status"] != "Pending":
            raise HTTPException(status_code=403, detail="You can only edit pending requests")
        # Students cannot change status, assign staff, or reject
        if update_data.status or update_data.assigned_staff_id or update_data.rejection_reason:
            raise HTTPException(status_code=403, detail="Students cannot update request status or assignments")
    
    now = datetime.now(timezone.utc).isoformat()
    updates = {"updated_at": now}
    
    old_status = request_doc["status"]
    
    if update_data.status:
        updates["status"] = update_data.status
        # Use the provided note or create a default one
        note_text = update_data.note if update_data.note else f"Status changed to {update_data.status}"
        timeline_entry = {
            "status": update_data.status,
            "timestamp": now,
            "note": note_text,
            "updated_by": current_user["full_name"]
        }
        await db.recommendation_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.assigned_staff_id:
        staff = await db.users.find_one({"id": update_data.assigned_staff_id}, {"_id": 0})
        if staff:
            updates["assigned_staff_id"] = update_data.assigned_staff_id
            updates["assigned_staff_name"] = staff["full_name"]
            # Notify staff
            await create_notification(
                staff["id"],
                "New Recommendation Assignment",
                f"You have been assigned a recommendation letter request",
                "recommendation_assignment",
                request_id
            )
    
    if update_data.rejection_reason:
        updates["rejection_reason"] = update_data.rejection_reason
        updates["status"] = "Rejected"
        timeline_entry = {
            "status": "Rejected",
            "timestamp": now,
            "note": f"Request rejected: {update_data.rejection_reason}",
            "updated_by": current_user["full_name"]
        }
        await db.recommendation_requests.update_one(
            {"id": request_id},
            {"$push": {"timeline": timeline_entry}}
        )
    
    if update_data.staff_notes:
        updates["staff_notes"] = update_data.staff_notes
    
    if update_data.co_curricular_activities is not None:
        updates["co_curricular_activities"] = update_data.co_curricular_activities
    
    await db.recommendation_requests.update_one({"id": request_id}, {"$set": updates})
    
    # Notify student of status change
    if update_data.status and update_data.status != old_status:
        student = await db.users.find_one({"id": request_doc["student_id"]}, {"_id": 0})
        if student:
            title = "Recommendation Request Status Updated"
            message = f"Your recommendation letter request has been updated from '{old_status}' to '{update_data.status}'."
            await create_notification(student["id"], title, message, "recommendation_status_update", request_id)
    
    updated_request = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    normalized_request = normalize_recommendation_data(updated_request)
    return RecommendationRequestResponse(**normalized_request)

@api_router.post("/recommendations/{request_id}/documents")
async def upload_recommendation_document(request_id: str, file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Only staff and admin can upload documents")
    
    request_doc = await db.recommendation_requests.find_one({"id": request_id}, {"_id": 0})
    if not request_doc:
        raise HTTPException(status_code=404, detail="Request not found")
    
    # Validate file type
    allowed_types = [
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "image/jpeg",
        "image/png",
        "image/gif"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not allowed")
    
    # Save file
    file_id = str(uuid.uuid4())
    file_ext = Path(file.filename).suffix
    file_path = UPLOAD_DIR / f"{file_id}{file_ext}"
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    now = datetime.now(timezone.utc).isoformat()
    doc_entry = {
        "id": file_id,
        "filename": file.filename,
        "content_type": file.content_type,
        "path": str(file_path),
        "uploaded_by": current_user["full_name"],
        "uploaded_at": now
    }
    
    await db.recommendation_requests.update_one(
        {"id": request_id},
        {
            "$push": {"documents": doc_entry},
            "$set": {"updated_at": now}
        }
    )
    
    # Add timeline entry
    timeline_entry = {
        "status": request_doc["status"],
        "timestamp": now,
        "note": f"Document uploaded: {file.filename}",
        "updated_by": current_user["full_name"]
    }
    await db.recommendation_requests.update_one(
        {"id": request_id},
        {"$push": {"timeline": timeline_entry}}
    )
    
    # Notify student
    await create_notification(
        request_doc["student_id"],
        "Document Uploaded",
        f"A document has been uploaded to your recommendation letter request",
        "recommendation_document",
        request_id
    )
    
    return {"message": "Document uploaded successfully", "document": doc_entry}

# ==================== NOTIFICATIONS ====================

@api_router.get("/notifications", response_model=List[NotificationResponse])
async def get_notifications(current_user: dict = Depends(get_current_user)):
    notifications = await db.notifications.find(
        {"user_id": current_user["id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    return [NotificationResponse(**n) for n in notifications]

@api_router.get("/notifications/unread-count")
async def get_unread_count(current_user: dict = Depends(get_current_user)):
    count = await db.notifications.count_documents({
        "user_id": current_user["id"],
        "read": False
    })
    return {"count": count}

@api_router.patch("/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.notifications.update_one(
        {"id": notification_id, "user_id": current_user["id"]},
        {"$set": {"read": True}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Notification not found")
    return {"message": "Notification marked as read"}

@api_router.patch("/notifications/read-all")
async def mark_all_notifications_read(current_user: dict = Depends(get_current_user)):
    await db.notifications.update_many(
        {"user_id": current_user["id"], "read": False},
        {"$set": {"read": True}}
    )
    return {"message": "All notifications marked as read"}

# ==================== ANALYTICS ====================

@api_router.get("/analytics", response_model=AnalyticsResponse)
async def get_analytics(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Only admins can view analytics")
    
    # Check and notify about overdue requests
    await check_and_notify_overdue_requests()
    
    now = datetime.now(timezone.utc)
    today_str = now.strftime("%Y-%m-%d")
    
    # Use aggregation pipeline for optimized queries
    pipeline = [
        {
            "$facet": {
                "status_counts": [
                    {"$group": {"_id": "$status", "count": {"$sum": 1}}}
                ],
                "enrollment_counts": [
                    {"$group": {"_id": "$enrollment_status", "count": {"$sum": 1}}}
                ],
                "collection_counts": [
                    {"$group": {"_id": "$collection_method", "count": {"$sum": 1}}}
                ],
                "total": [
                    {"$count": "count"}
                ],
                "staff_workload": [
                    {"$match": {"assigned_staff_id": {"$ne": None}}},
                    {"$group": {"_id": "$assigned_staff_id", "count": {"$sum": 1}}}
                ]
            }
        }
    ]
    
    result = await db.transcript_requests.aggregate(pipeline).to_list(1)
    
    # Calculate overdue requests (needed_by_date < today and status not Completed/Rejected)
    overdue_requests = await db.transcript_requests.find({
        "needed_by_date": {"$lt": today_str},
        "status": {"$nin": ["Completed", "Rejected"]}
    }).to_list(None)
    overdue_count = len(overdue_requests)
    
    # Calculate overdue by days categories
    overdue_by_days = {"1-3 days": 0, "4-7 days": 0, "8-14 days": 0, "15+ days": 0}
    for req in overdue_requests:
        try:
            needed_date = datetime.strptime(req["needed_by_date"], "%Y-%m-%d")
            days_overdue = (now.replace(tzinfo=None) - needed_date).days
            if days_overdue <= 3:
                overdue_by_days["1-3 days"] += 1
            elif days_overdue <= 7:
                overdue_by_days["4-7 days"] += 1
            elif days_overdue <= 14:
                overdue_by_days["8-14 days"] += 1
            else:
                overdue_by_days["15+ days"] += 1
        except:
            pass
    
    overdue_by_days_list = [{"name": k, "value": v, "color": "#ef4444" if "15+" in k else "#f97316" if "8-14" in k else "#eab308" if "4-7" in k else "#fbbf24"} for k, v in overdue_by_days.items() if v > 0]
    
    # Parse aggregation results
    if result:
        data = result[0]
        
        # Get total
        total = data["total"][0]["count"] if data["total"] else 0
        
        # Parse status counts
        status_map = {item["_id"]: item["count"] for item in data["status_counts"]}
        pending = status_map.get("Pending", 0)
        in_progress = status_map.get("In Progress", 0)
        processing = status_map.get("Processing", 0)
        ready = status_map.get("Ready", 0)
        completed = status_map.get("Completed", 0)
        rejected = status_map.get("Rejected", 0)
        
        # Parse enrollment counts
        enrollment_map = {item["_id"]: item["count"] for item in data["enrollment_counts"]}
        requests_by_enrollment = [
            {"name": "Enrolled", "value": enrollment_map.get("enrolled", 0)},
            {"name": "Graduate", "value": enrollment_map.get("graduate", 0)},
            {"name": "Withdrawn", "value": enrollment_map.get("withdrawn", 0)}
        ]
        
        # Parse collection method counts
        collection_map = {item["_id"]: item["count"] for item in data["collection_counts"]}
        requests_by_collection_method = [
            {"name": "Pickup at Bursary", "value": collection_map.get("pickup", 0)},
            {"name": "Emailed to Institution", "value": collection_map.get("emailed", 0)},
            {"name": "Physical Delivery", "value": collection_map.get("delivery", 0)}
        ]
        
        # Parse staff workload
        staff_workload_map = {item["_id"]: item["count"] for item in data["staff_workload"]}
        staff_workload = []
        for staff_id, count in staff_workload_map.items():
            staff = await db.users.find_one({"id": staff_id})
            staff_name = staff["full_name"] if staff else "Unknown"
            staff_workload.append({"name": staff_name, "requests": count})
        
        # Add unassigned count
        unassigned_count = await db.transcript_requests.count_documents({
            "$or": [{"assigned_staff_id": None}, {"assigned_staff_id": {"$exists": False}}]
        })
        if unassigned_count > 0:
            staff_workload.append({"name": "Unassigned", "requests": unassigned_count})
        
    else:
        total = pending = in_progress = processing = ready = completed = rejected = 0
        requests_by_enrollment = [
            {"name": "Enrolled", "value": 0},
            {"name": "Graduate", "value": 0},
            {"name": "Withdrawn", "value": 0}
        ]
        requests_by_collection_method = [
            {"name": "Pickup at Bursary", "value": 0},
            {"name": "Emailed to Institution", "value": 0},
            {"name": "Physical Delivery", "value": 0}
        ]
        staff_workload = []
    
    # Requests by month (last 6 months) - using aggregation
    requests_by_month = []
    
    for i in range(5, -1, -1):
        month_start = (now.replace(day=1) - timedelta(days=i*30)).replace(day=1)
        month_end = (month_start + timedelta(days=32)).replace(day=1)
        
        count = await db.transcript_requests.count_documents({
            "created_at": {
                "$gte": month_start.isoformat(),
                "$lt": month_end.isoformat()
            }
        })
        requests_by_month.append({
            "month": month_start.strftime("%b %Y"),
            "count": count
        })
    
    # Get recommendation letter stats
    total_rec = await db.recommendation_requests.count_documents({})
    pending_rec = await db.recommendation_requests.count_documents({"status": "Pending"})
    completed_rec = await db.recommendation_requests.count_documents({"status": "Completed"})
    in_progress_rec = await db.recommendation_requests.count_documents({"status": "In Progress"})
    rejected_rec = await db.recommendation_requests.count_documents({"status": "Rejected"})
    
    # Get overdue recommendation requests
    overdue_rec_count = 0
    overdue_rec_by_days = []
    rec_requests = await db.recommendation_requests.find({
        "status": {"$nin": ["Completed", "Rejected"]},
        "needed_by_date": {"$ne": None, "$ne": ""}
    }, {"_id": 0, "needed_by_date": 1}).to_list(10000)
    
    for req in rec_requests:
        try:
            needed_date = datetime.fromisoformat(req["needed_by_date"].replace('Z', '+00:00')).date()
            if needed_date < now.date():
                overdue_rec_count += 1
                days_overdue = (now.date() - needed_date).days
                if days_overdue <= 7:
                    overdue_rec_by_days.append({"days": "1-7 days", "count": 1})
                elif days_overdue <= 14:
                    overdue_rec_by_days.append({"days": "8-14 days", "count": 1})
                elif days_overdue <= 30:
                    overdue_rec_by_days.append({"days": "15-30 days", "count": 1})
                else:
                    overdue_rec_by_days.append({"days": "30+ days", "count": 1})
        except:
            pass
    
    # Aggregate overdue recommendation by days
    rec_overdue_agg = {}
    for item in overdue_rec_by_days:
        rec_overdue_agg[item["days"]] = rec_overdue_agg.get(item["days"], 0) + 1
    overdue_rec_by_days_list = [{"days": k, "count": v} for k, v in rec_overdue_agg.items()]
    
    # Get recommendation collection method breakdown
    rec_collection_counts = await db.recommendation_requests.aggregate([
        {"$group": {"_id": "$collection_method", "count": {"$sum": 1}}}
    ]).to_list(10)
    rec_collection_map = {item["_id"]: item["count"] for item in rec_collection_counts if item["_id"]}
    recommendations_by_collection_method = [
        {"name": "Pickup at School", "value": rec_collection_map.get("pickup", 0)},
        {"name": "Emailed to Institution", "value": rec_collection_map.get("emailed", 0)},
        {"name": "Physical Delivery", "value": rec_collection_map.get("delivery", 0)}
    ]
    
    return AnalyticsResponse(
        total_requests=total,
        pending_requests=pending,
        in_progress_requests=in_progress,
        processing_requests=processing,
        ready_requests=ready,
        completed_requests=completed,
        rejected_requests=rejected,
        overdue_requests=overdue_count,
        requests_by_month=requests_by_month,
        requests_by_enrollment=requests_by_enrollment,
        requests_by_collection_method=requests_by_collection_method,
        staff_workload=staff_workload,
        overdue_by_days=overdue_by_days_list,
        total_recommendation_requests=total_rec,
        pending_recommendation_requests=pending_rec,
        in_progress_recommendation_requests=in_progress_rec,
        completed_recommendation_requests=completed_rec,
        rejected_recommendation_requests=rejected_rec,
        overdue_recommendation_requests=overdue_rec_count,
        recommendations_by_collection_method=recommendations_by_collection_method,
        overdue_transcripts_by_days=overdue_by_days_list,
        overdue_recommendations_by_days=overdue_rec_by_days_list
    )

# ==================== HEALTH CHECK ====================

@api_router.get("/")
async def root():
    return {"message": "WBS Transcript and Recommendation Tracker API", "status": "running"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy"}

# ==================== EXPORT/REPORTING ENDPOINTS ====================

def format_date_for_export(date_str):
    """Format date string for export"""
    try:
        if date_str:
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return dt.strftime("%Y-%m-%d %H:%M")
    except:
        pass
    return date_str or ""

def format_years_for_export(years_data):
    """Format years attended/academic years for export"""
    if isinstance(years_data, list):
        return ", ".join([f"{y.get('from_year', '')}-{y.get('to_year', '')}" for y in years_data])
    return str(years_data) if years_data else ""

@api_router.get("/export/transcripts/{format_type}")
async def export_transcript_requests(format_type: str, status: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """Export transcript requests as DOCX, PDF, or XLSX"""
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    # Build query
    query = {}
    if current_user["role"] == "staff":
        query["assigned_staff_id"] = current_user["id"]
    if status and status != "all":
        query["status"] = status
    
    requests = await db.transcript_requests.find(query, {"_id": 0}).sort("created_at", -1).to_list(10000)
    
    if format_type == "xlsx":
        return generate_transcript_xlsx(requests)
    elif format_type == "pdf":
        return generate_transcript_pdf(requests)
    elif format_type == "docx":
        return generate_transcript_docx(requests)
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use xlsx, pdf, or docx")

@api_router.get("/export/recommendations/{format_type}")
async def export_recommendation_requests(format_type: str, status: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """Export recommendation requests as DOCX, PDF, or XLSX"""
    if current_user["role"] not in ["admin", "staff"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    
    # Build query
    query = {}
    if current_user["role"] == "staff":
        query["assigned_staff_id"] = current_user["id"]
    if status and status != "all":
        query["status"] = status
    
    requests = await db.recommendation_requests.find(query, {"_id": 0}).sort("created_at", -1).to_list(10000)
    
    if format_type == "xlsx":
        return generate_recommendation_xlsx(requests)
    elif format_type == "pdf":
        return generate_recommendation_pdf(requests)
    elif format_type == "docx":
        return generate_recommendation_docx(requests)
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use xlsx, pdf, or docx")

def generate_transcript_xlsx(requests):
    """Generate XLSX file for transcript requests"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Transcript Requests"
    
    # Headers
    headers = ["ID", "Student Name", "Email", "School ID", "Status", "Academic Years", 
               "Collection Method", "Institution", "Needed By", "Assigned Staff", "Created At"]
    
    # Style headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="800000", end_color="800000", fill_type="solid")
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for row_num, req in enumerate(requests, 2):
        ws.cell(row=row_num, column=1, value=req.get("id", "")[:8])
        ws.cell(row=row_num, column=2, value=req.get("student_name", ""))
        ws.cell(row=row_num, column=3, value=req.get("student_email", ""))
        ws.cell(row=row_num, column=4, value=req.get("school_id", ""))
        ws.cell(row=row_num, column=5, value=req.get("status", ""))
        ws.cell(row=row_num, column=6, value=format_years_for_export(req.get("academic_years", req.get("academic_year", ""))))
        ws.cell(row=row_num, column=7, value=req.get("collection_method", ""))
        ws.cell(row=row_num, column=8, value=req.get("institution_name", ""))
        ws.cell(row=row_num, column=9, value=format_date_for_export(req.get("needed_by_date", "")))
        ws.cell(row=row_num, column=10, value=req.get("assigned_staff_name", "Unassigned"))
        ws.cell(row=row_num, column=11, value=format_date_for_export(req.get("created_at", "")))
    
    # Adjust column widths
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 40)
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=transcript_requests_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

def generate_transcript_pdf(requests):
    """Generate PDF file for transcript requests"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=30, bottomMargin=30)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, spaceAfter=20)
    elements.append(Paragraph("Transcript Requests Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Table data
    data = [["ID", "Student", "Status", "Academic Years", "Collection", "Institution", "Needed By", "Staff"]]
    
    for req in requests:
        data.append([
            req.get("id", "")[:8],
            req.get("student_name", ""),
            req.get("status", ""),
            format_years_for_export(req.get("academic_years", req.get("academic_year", "")))[:20],
            req.get("collection_method", ""),
            (req.get("institution_name", "") or "")[:20],
            format_date_for_export(req.get("needed_by_date", ""))[:10],
            (req.get("assigned_staff_name", "") or "Unassigned")[:15]
        ])
    
    # Create table
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.5, 0, 0)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=transcript_requests_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

def generate_transcript_docx(requests):
    """Generate DOCX file for transcript requests"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Transcript Requests Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total Requests: {len(requests)}")
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["Student", "Status", "Academic Years", "Collection", "Institution", "Needed By", "Staff"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for req in requests:
        row_cells = table.add_row().cells
        row_cells[0].text = req.get("student_name", "")
        row_cells[1].text = req.get("status", "")
        row_cells[2].text = format_years_for_export(req.get("academic_years", req.get("academic_year", "")))
        row_cells[3].text = req.get("collection_method", "")
        row_cells[4].text = req.get("institution_name", "") or ""
        row_cells[5].text = format_date_for_export(req.get("needed_by_date", ""))[:10]
        row_cells[6].text = req.get("assigned_staff_name", "") or "Unassigned"
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=transcript_requests_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

def generate_recommendation_xlsx(requests):
    """Generate XLSX file for recommendation requests"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Recommendation Requests"
    
    # Headers
    headers = ["ID", "Student Name", "Email", "Status", "Years Attended", "Form Class",
               "Institution", "Program", "Collection Method", "Needed By", "Assigned Staff", "Created At"]
    
    # Style headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="DAA520", end_color="DAA520", fill_type="solid")
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for row_num, req in enumerate(requests, 2):
        ws.cell(row=row_num, column=1, value=req.get("id", "")[:8])
        ws.cell(row=row_num, column=2, value=req.get("student_name", ""))
        ws.cell(row=row_num, column=3, value=req.get("student_email", ""))
        ws.cell(row=row_num, column=4, value=req.get("status", ""))
        ws.cell(row=row_num, column=5, value=format_years_for_export(req.get("years_attended", req.get("years_attended_str", ""))))
        ws.cell(row=row_num, column=6, value=req.get("last_form_class", ""))
        ws.cell(row=row_num, column=7, value=req.get("institution_name", ""))
        ws.cell(row=row_num, column=8, value=req.get("program_name", ""))
        ws.cell(row=row_num, column=9, value=req.get("collection_method", ""))
        ws.cell(row=row_num, column=10, value=format_date_for_export(req.get("needed_by_date", "")))
        ws.cell(row=row_num, column=11, value=req.get("assigned_staff_name", "Unassigned"))
        ws.cell(row=row_num, column=12, value=format_date_for_export(req.get("created_at", "")))
    
    # Adjust column widths
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 40)
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=recommendation_requests_{datetime.now().strftime('%Y%m%d')}.xlsx"}
    )

def generate_recommendation_pdf(requests):
    """Generate PDF file for recommendation requests"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=30, bottomMargin=30)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, spaceAfter=20)
    elements.append(Paragraph("Recommendation Letter Requests Report", title_style))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 20))
    
    # Table data
    data = [["ID", "Student", "Status", "Years", "Institution", "Program", "Collection", "Needed By", "Staff"]]
    
    for req in requests:
        data.append([
            req.get("id", "")[:8],
            req.get("student_name", ""),
            req.get("status", ""),
            format_years_for_export(req.get("years_attended", req.get("years_attended_str", "")))[:15],
            (req.get("institution_name", "") or "")[:18],
            (req.get("program_name", "") or "")[:18],
            req.get("collection_method", ""),
            format_date_for_export(req.get("needed_by_date", ""))[:10],
            (req.get("assigned_staff_name", "") or "Unassigned")[:12]
        ])
    
    # Create table
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.85, 0.65, 0.13)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightyellow),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 7),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=recommendation_requests_{datetime.now().strftime('%Y%m%d')}.pdf"}
    )

def generate_recommendation_docx(requests):
    """Generate DOCX file for recommendation requests"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Recommendation Letter Requests Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total Requests: {len(requests)}")
    doc.add_paragraph()
    
    # Create table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Headers
    headers = ["Student", "Status", "Years", "Form Class", "Institution", "Program", "Needed By", "Staff"]
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for req in requests:
        row_cells = table.add_row().cells
        row_cells[0].text = req.get("student_name", "")
        row_cells[1].text = req.get("status", "")
        row_cells[2].text = format_years_for_export(req.get("years_attended", req.get("years_attended_str", "")))
        row_cells[3].text = req.get("last_form_class", "")
        row_cells[4].text = req.get("institution_name", "") or ""
        row_cells[5].text = req.get("program_name", "") or ""
        row_cells[6].text = format_date_for_export(req.get("needed_by_date", ""))[:10]
        row_cells[7].text = req.get("assigned_staff_name", "") or "Unassigned"
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=recommendation_requests_{datetime.now().strftime('%Y%m%d')}.docx"}
    )

# ==================== SEED DEFAULT ADMIN ====================

@app.on_event("startup")
async def seed_default_admin():
    admin_email = "admin@wolmers.org"
    existing = await db.users.find_one({"email": admin_email}, {"_id": 0})
    
    if not existing:
        admin_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()
        
        admin_doc = {
            "id": admin_id,
            "email": admin_email,
            "full_name": "System Administrator",
            "password_hash": hash_password("Admin123!"),
            "role": "admin",
            "created_at": now,
            "updated_at": now
        }
        
        await db.users.insert_one(admin_doc)
        logger.info("Default admin account created: admin@wolmers.org / Admin123!")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
